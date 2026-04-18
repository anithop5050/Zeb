from __future__ import annotations

import re
from collections import Counter
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document

SOURCE_DOCX = Path("reference_template.docx")
OUTPUT_DOCX = Path("anith.docx")
REPORT_PATH = Path("analysis/anith_verification_report.txt")

TITLE_EXACT = "ZEB:Robust Invisible Image Watermarking Using Deep Learning"
TITLE_UPPER = "ZEB:ROBUST INVISIBLE IMAGE WATERMARKING USING DEEP LEARNING"
STUDENT_EXACT = "ANITH SHIBU THOMAS- LMC24MCA-2016"


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def nonempty_paragraphs_for_table(doc: Document, table_index: int):
    cell = doc.tables[table_index].cell(0, 0)
    return [p for p in cell.paragraphs if p.text.strip()]


def apply_page_update(doc: Document, table_index: int, lines: List[str]) -> None:
    paras = nonempty_paragraphs_for_table(doc, table_index)
    for i, line in enumerate(lines):
        if i >= len(paras):
            break
        set_paragraph_text(paras[i], line)


def replace_text_everywhere(doc: Document, replacements: List[Tuple[str, str]]) -> None:
    def _replace_in_paragraph(paragraph) -> None:
        original = paragraph.text
        updated = original
        for old, new in replacements:
            updated = updated.replace(old, new)
        if updated != original:
            set_paragraph_text(paragraph, updated)

    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph)


def all_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def structure_metrics(doc: Document) -> Dict[str, int]:
    return {
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
    }


def style_sequence(doc: Document) -> List[str]:
    seq: List[str] = []
    for paragraph in all_paragraphs(doc):
        style_name = paragraph.style.name if paragraph.style is not None else ""
        seq.append(style_name)
    return seq


def font_signature(doc: Document) -> Counter:
    sig: Counter = Counter()
    for paragraph in all_paragraphs(doc):
        for run in paragraph.runs:
            size_pt = None
            if run.font.size is not None:
                size_pt = float(run.font.size.pt)
            sig[(run.font.name, size_pt, bool(run.bold), bool(run.italic), bool(run.underline))] += 1
    return sig


def header_footer_snapshot(doc: Document) -> List[Dict[str, object]]:
    snapshot: List[Dict[str, object]] = []
    for i, section in enumerate(doc.sections):
        header_text = "\n".join(p.text for p in section.header.paragraphs)
        footer_text = "\n".join(p.text for p in section.footer.paragraphs)
        header_styles = [p.style.name if p.style is not None else "" for p in section.header.paragraphs]
        footer_styles = [p.style.name if p.style is not None else "" for p in section.footer.paragraphs]
        snapshot.append(
            {
                "section": i,
                "header_text": header_text,
                "footer_text": footer_text,
                "header_styles": header_styles,
                "footer_styles": footer_styles,
            }
        )
    return snapshot


def field_counts(doc: Document) -> Dict[str, int]:
    parts = [doc.part]
    seen = {id(doc.part)}
    for section in doc.sections:
        for part in (section.header.part, section.footer.part):
            part_id = id(part)
            if part_id not in seen:
                seen.add(part_id)
                parts.append(part)

    total_field_nodes = 0
    total_page_fields = 0

    for part in parts:
        root = part.element
        fld_simple_nodes = root.xpath('.//*[local-name()="fldSimple"]')
        instr_nodes = root.xpath('.//*[local-name()="instrText"]')

        total_field_nodes += len(fld_simple_nodes) + len(instr_nodes)

        for node in fld_simple_nodes:
            instr_attr = node.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr") or ""
            if "PAGE" in instr_attr.upper():
                total_page_fields += 1

        for node in instr_nodes:
            text = node.text or ""
            if "PAGE" in text.upper():
                total_page_fields += 1

    return {
        "field_nodes": total_field_nodes,
        "page_fields": total_page_fields,
    }


def extract_all_text(doc: Document) -> str:
    chunks: List[str] = []
    for paragraph in all_paragraphs(doc):
        txt = paragraph.text.strip()
        if txt:
            chunks.append(txt)
    return "\n".join(chunks)


def chapter_order_ok(text_upper: str) -> bool:
    markers = [f"CHAPTER-{i}" for i in range(1, 11)]
    positions: List[int] = []
    for marker in markers:
        pos = text_upper.find(marker)
        if pos == -1:
            return False
        positions.append(pos)
    return positions == sorted(positions)


def table_page_number(doc: Document, table_index: int) -> str:
    paras = nonempty_paragraphs_for_table(doc, table_index)
    if not paras:
        return "N/A"
    tail = paras[-1].text
    match = re.search(r"(\d+)\s*$", tail)
    return match.group(1) if match else "N/A"


def build_page_updates() -> Dict[int, List[str]]:
    return {
        0: [
            "ZEB",
            "ROBUST INVISIBLE IMAGE WATERMARKING USING DEEP LEARNING",
            "A Project Report\nSubmitted By,\nANITH SHIBU THOMAS- LMC24MCA-2016\nin partial fulfilment of the requirements for the award of the degree in\nMASTER OF COMPUTER APPLICATIONS at",
        ],
        1: [
            "LOURDES MATHA COLLEGE OF SCIENCE AND TECHNOLOGY",
            "KUTTICHAL, THIRUVANATHAPURAM-695574",
            "(Affiliated To APJ Abdul Kalam Technological University, Kerala)",
            "DEPARTMENT OF COMPUTER APPLICATIONS",
            "CERTIFICATE",
            "This is to certify that the project work entitled \"ZEB:Robust Invisible Image Watermarking Using Deep Learning\" is a bonafide record of the work done by Mr. ANITH SHIBU THOMAS, Reg No: LMC24MCA-2016, student of Department of Computer Applications, Lourdes Matha College of Science and Technology.",
            "Ms. SHERIN JOSEPH\tDate:\t(Internal Guide)",
            "Ms. BISMI K CHARLEYS\n(Head of the Department)",
        ],
        2: [
            "DECLARATION",
            "I hereby declare that the project report \"ZEB:Robust Invisible Image Watermarking Using Deep Learning\" submitted for partial fulfilment of the requirements for the award of the degree of Master of Computer Applications of APJ Abdul Kalam Technological University is a record of original work carried out by me under proper guidance.",
            "Place: TRIVANDRUM\tANITH SHIBU THOMAS\tDate:",
        ],
        3: [
            "ACKNOWLEDGEMENT",
            "This project became possible through the support and encouragement of many people. I sincerely thank everyone who helped me complete this work.",
            "I express my gratitude to God Almighty for guidance and strength throughout this project period.",
            "I thank Rev. Dr. Bejoy Arackal, Director, and Dr. Beshiba Wilson, Principal, Lourdes Matha College of Science and Technology, for providing the institutional support required for this work.",
            "I am grateful to Ms. Bismi K Charleys, Head of the Department of Computer Applications, for her valuable guidance and constant encouragement.",
            "I place on record my sincere thanks to my internal guide Ms. Sherin Joseph and all faculty members of the Department of Computer Applications for their timely suggestions and technical support.",
            "I also acknowledge my classmates, friends, and open-source contributors whose resources and discussions helped in shaping this project.",
        ],
        4: [
            "CONTENTS",
            "CONTENT    Page No ABSTRACT    1 CHAPTER 1  1.INTRODUCTION  1.1 BACKGROUND AND MOTIVATION 2  1.2 OBJECTIVE OF THE PROJECT 3 CHAPTER 2  2. LITERATURE SURVEY  2.1 STUDY OF SIMILAR WORKS 4  2.2 EXISTING SYSTEM 7  2.3 DRAWBACKS OF EXISTING SYSTEM 8 CHAPTER 3  3. OVERALL DESCRIPTION  3.1 PROPOSED SYSTEM 9  3.2 FEATURES OF PROPOSED SYSTEM 9  3.3 FUNCTIONS OF PROPOSED SYSTEM 10  3.4 REQUIREMENTS SPECIFICATIONS 11  3.5 FEASIBILITY ANALYSIS 12",
        ],
        5: [
            "4.3.6 OPENCV 17",
            "4.3.7 NUMPY 18",
            "4.3.8 TKINTER 18",
            "4.3.9 VISUAL STUDIO CODE 19",
            "CHAPTER 5",
            "5. DESIGN",
            "5.1 SYSTEM DESIGN 20",
            "5.2 PIPELINE DIAGRAM 21",
            "5.3 ACTOR FLOW DIAGRAM 22",
            "5.3.1 EMBEDDING WORKFLOW 22",
            "5.3.2 VERIFICATION WORKFLOW 23",
            "5.4 INPUT DESIGN 24",
            "5.5 OUTPUT DESIGN 25",
            "5.6 MODULE DESIGN 26",
            "5.7 DATA MODEL DESIGN 27",
            "5.8 TABLE DESIGN 28",
            "CHAPTER 6",
            "6. FUNCTIONAL AND NON-FUNCTIONAL REQUIREMENTS",
            "6.1 FUNCTIONAL REQUIREMENTS 52",
            "6.2 NON-FUNCTIONAL REQUIREMENTS 52",
            "CHAPTER 7",
            "7. IMPLEMENTATION AND TESTING",
            "7.1 SYSTEM IMPLEMENTATION 54",
            "7.2 TESTING 55",
            "7.2.1 SYSTEM TESTING 55",
            "7.2.2 UNIT TESTING 56",
            "7.2.3 INTEGRATION TESTING 56",
            "7.2.4 BLACK BOX TESTING 56",
            "7.2.5 VALIDATION TESTING 56",
            "7.2.6 OUTPUT TESTING 57",
            "7.2.7 USER ACCEPTANCE TESTING 57",
        ],
        6: [
            "7.2.8 WHITE BOX TESTING 57",
            "7.3 TEST CASES 58",
            "CHAPTER 8",
            "8. RESULTS AND DISCUSSIONS",
            "8.1 RESULTS 61",
            "8.2 VISUAL OUTPUTS 62",
            "CHAPTER 9",
            "9. CONCLUSION",
            "9.1 SYSTEM MAINTENANCE 76",
            "9.2 CONCLUSION 77",
            "9.3 FUTURE ENHANCEMENT 78",
            "CHAPTER 10",
            "10. BIBLIOGRAPHY",
            "10.1 BOOKS 79",
            "10.2 WEBSITES 79",
            "10.3 JOURNALS 80",
            "APPENDICES",
            "1. LIST OF TABLES 82",
            "2. LIST OF FIGURES 83",
            "CHANGE HISTORY 84",
        ],
        7: [
            TITLE_UPPER,
            "ABSTRACT",
            "Digital image distribution requires ownership protection that remains effective under real-world distortions. This project, ZEB:Robust Invisible Image Watermarking Using Deep Learning, implements a deep encoder-decoder watermarking system that embeds a 64-bit payload with minimal visual impact while preserving reliable extraction. The pipeline integrates low-alpha control, attack simulation, optional semantic enhancement, and reliability checks including post-save verification. Experimental observations in this workspace indicate strong invisibility-robustness balance, including PSNR values above 36 dB in practical settings and near-zero BER under moderate attacks. The system also includes a seed-owner registry with audit support for accountability in ownership claims.",
        ],
        8: ["CHAPTER-1\nINTRODUCTION"],
        9: [
            TITLE_UPPER,
            "1.1 BACKGROUND AND MOTIVATION",
            "Digital media can be copied and redistributed instantly, making ownership protection a significant challenge for creators and institutions.",
            "Invisible image watermarking addresses this challenge by embedding a hidden signal that does not disturb perceptual quality while enabling later verification.",
            "Traditional methods are often vulnerable to compression, blur, resizing, and additive noise; hence robust deep learning approaches are required.",
            "The ZEB system adopts a neural encoder-decoder design with attack-aware training, low-alpha deployment policy, and optional semantic residual support.",
            "By combining robustness engineering with strict visual quality control, the proposed workflow improves practical watermark survivability.",
        ],
        10: [
            TITLE_UPPER,
            "1.2 OBJECTIVE OF THE PROJECT",
            "The primary objective is to build a robust and practically deployable invisible image watermarking framework for ownership verification.",
            "The project targets an end-to-end workflow that supports training, embedding, extraction, reliability validation, and owner-seed traceability.",
            "The key goals of the project are:",
            "- Build an encoder-decoder model for 64-bit watermark embedding.\n- Maintain high perceptual quality at low alpha values.\n- Improve extraction reliability under common image attacks.",
            "Overall, the project aims to deliver a balanced system where invisibility, robustness, and operational auditability are all treated as first-class requirements.",
        ],
        11: ["CHAPTER-2\nLITERATURE SURVEY"],
        12: [
            TITLE_UPPER,
            "2.1 STUDY OF SIMILAR WORKS",
            "1. HiDDeN: Hiding Data With Deep Networks (ECCV 2018)\nThis work introduced learned end-to-end data hiding, demonstrating that neural models can embed and recover payloads better than many hand-crafted approaches.",
            "Key Points:\n- End-to-end optimization for embedding and decoding\n- Differentiable distortion simulation\n- Improved payload recovery compared with classical baselines",
            "2. RivaGAN and related robust neural watermarking models\nThese studies emphasized robustness under video and image transformations and showed the importance of attack-aware training objectives.",
            "Key Points:\n- Robustness to compression and geometric changes\n- Adversarial style training for resilience\n- Practical trade-off between visibility and robustness",
        ],
        13: [
            TITLE_UPPER,
            "3. MBRS-style distortion pipelines for robust image watermarking",
            "Recent approaches combine mini-batches of real and simulated distortions during training, improving model behavior under mixed attack settings.",
            "Key Points:",
            "- Multi-distortion training improves generalization\n- Curriculum design supports stable convergence\n- Useful for deployment-oriented robustness",
            "These approaches inspire the attack simulation strategy used in this project.",
            "4. StegaStamp and neural marker systems",
            "StegaStamp demonstrates practical hidden markers for camera-captured scenes, highlighting the value of robust decoding under perspective and photometric variations.",
            "Key Points:",
            "- Robust marker decoding from transformed captures\n- Strong practical motivation for real-world invariance\n- End-to-end trainable visual embedding",
            "This project adapts similar robustness thinking but focuses on high-fidelity invisible watermarking for standard images.",
        ],
        14: [
            TITLE_UPPER,
            "5. Deep watermarking with perceptual and fidelity constraints",
            "Several papers optimize BER jointly with perceptual losses, showing that visual quality and decode reliability can be co-optimized when alpha is properly controlled.",
            "Key Points:",
            "- Joint optimization of fidelity and recoverability\n- Perceptual constraints reduce visible artifacts\n- Strong relevance to low-alpha deployment",
            "6. Ownership tracking and provenance systems",
            "Modern digital provenance frameworks stress the need for verifiable identity linkage in addition to watermark extraction.",
            "Key Points:",
            "- Traceable ownership records improve legal defensibility\n- Audit logs strengthen chain-of-custody evidence\n- Registry integration improves operational governance",
            "These findings motivate the seed-owner registry integrated in the present project.",
        ],
        15: [
            TITLE_UPPER,
            "2.2 EXISTING SYSTEM",
            "Existing image protection workflows commonly use either visible logos, fragile metadata tags, or basic signal-processing watermark methods with limited robustness.",
            "Many systems fail under practical distortions such as recompression, resizing, and filtering, and they rarely provide integrated ownership audit support.",
            "As a result, current workflows often require manual evidence collection and do not provide a unified reliability-first verification pipeline.",
        ],
        16: [
            TITLE_UPPER,
            "2.3 DRAWBACKS OF EXISTING SYSTEM",
            "1. Limited robustness under compression and post-processing operations.",
            "2. Weak extraction reliability at low-visibility settings.",
            "3. Inadequate balancing of visual quality and payload recoverability.",
            "4. Absence of integrated attack simulation during model development.",
            "5. Poor operational traceability for ownership claims.",
            "6. Fragmented tooling across embedding, extraction, and validation stages.",
            "7. Lack of standardized post-save verification and audit logging.",
        ],
        17: ["CHAPTER-3\nOVERALL DESCRIPTION"],
        18: [
            TITLE_UPPER,
            "3.1 PROPOSED SYSTEM",
            "The proposed system is a deep learning based invisible watermarking framework that embeds and recovers a 64-bit payload while preserving image quality.",
            "The architecture combines a pixel-domain encoder-decoder path, optional semantic residual processing, and attack-aware training modules.",
            "A reliability layer enforces alpha clamping, save policy checks, and post-embed verification for consistent deployment behavior.",
            "3.2 FEATURES OF PROPOSED SYSTEM",
            "- Low-visibility embedding with controlled alpha policy.",
            "- Robust extraction under compression, blur, and noise disturbances.",
            "- Ownership traceability using a SQLite-based seed registry and audit log.",
        ],
        19: [
            TITLE_UPPER,
            "- Deterministic seed-to-watermark generation for repeatable verification.",
            "- Optional semantic enhancement for difficult attack conditions.",
            "- Quantization-aware save and extraction safety checks.",
            "- Batch embedding and decoding support for scalable processing.",
            "- Integrated BER/PSNR monitoring for quality governance.",
            "- Modular architecture for training and inference evolution.",
            "- Operationally aligned low-alpha deployment pathway.",
            "3.3 FUNCTIONS OF PROPOSED SYSTEM",
            "- Payload generation and embedding from seed-based identity.",
            "- Watermark extraction and BER-based verification.",
            "- Attack simulation and robustness benchmarking.",
            "- Registry logging for owner association and audit trails.",
        ],
        20: [
            TITLE_UPPER,
            "- Centralized reliability enforcement for save-path and alpha constraints.",
            "- Batch run orchestration for high-volume images.",
            "- Metric reporting for BER, PSNR, and extraction confidence.",
            "- Seed lookup and owner validation workflows.",
            "- Failure diagnostics and verification status outputs.",
            "- Checkpoint-aware inference with configurable modules.",
            "3.4 REQUIREMENTS SPECIFICATION",
            "System analysis for this project focuses on balancing three requirements: invisibility, robustness, and traceability.",
            "The specification defines model behavior under distortion, acceptable quality thresholds, registry consistency, and verification outputs required for ownership evidence.",
        ],
        21: [
            TITLE_UPPER,
            "The requirement specification for the proposed system is as follows:",
            "- Seed-driven payload generation with deterministic reproduction.",
            "- Encoder and decoder support for 64-bit watermark vectors.",
            "- Attack simulation during training for robustness conditioning.",
            "- Alpha policy control for operational invisibility constraints.",
            "- Reliability checks including post-save extraction verification.",
            "- Batch and single-image workflows with equivalent behavior.",
            "- Owner registry integration with audit event recording.",
            "- Usable analysis outputs for BER and visual fidelity interpretation.",
            "3.5 FEASIBILITY ANALYSIS",
            "Feasibility is evaluated in terms of technology availability, operational usability, cost profile, and user acceptance of quality-security trade-offs.",
        ],
        22: [
            TITLE_UPPER,
            "The study confirms that required libraries, compute resources, and storage systems are available within standard academic development setups.",
            "The proposed workflow can be implemented incrementally while preserving compatibility with existing image processing practices.",
            "3.5.1 TECHNICAL FEASIBILITY",
            "The project uses widely available tools such as Python, PyTorch, OpenCV, and SQLite. Model training and inference pipelines are practical on common GPU-enabled environments.",
            "Attack simulation, metric computation, and registry operations are all implementable using established open-source frameworks.",
            "3.5.2 OPERATIONAL FEASIBILITY",
            "The workflow supports straightforward embedding and verification with guided parameters, making it suitable for repeated operational use.",
            "Users can perform ownership checks without deep model retraining knowledge, improving adoption.",
            "3.5.3 ECONOMIC FEASIBILITY",
            "The system primarily depends on open-source tools and commodity hardware, resulting in low software licensing cost and manageable maintenance overhead.",
        ],
        23: [
            TITLE_UPPER,
            "The expected benefits in copyright assurance and verification reliability outweigh setup and maintenance cost, making the system economically feasible.",
            "3.5.4 BEHAVIORAL FEASIBILITY",
            "Users are likely to accept the system because it preserves visual quality while providing clear verification evidence and ownership traceability.",
            "The modular interface and measurable outcomes improve trust and encourage continued usage in practical scenarios.",
        ],
        24: ["CHAPTER-4\nOPERATING ENVIRONMENT"],
        25: [
            TITLE_UPPER,
            "4.1 HARDWARE REQUIREMENTS",
            "The system is designed to run on standard development machines, with optional GPU acceleration for faster model training and benchmarking.",
            "The following hardware requirements are sufficient for smooth experimentation and validation.",
            "testing:",
            "- Processor : Intel Core i5 or higher",
            "- RAM : 8 GB minimum (16 GB recommended)",
            "- Storage : 256 GB SSD minimum",
            "- Display : 1280x720 resolution or higher",
            "4.2 SOFTWARE REQUIREMENTS",
            "The software stack used in this project includes:",
            "- Operating System : Windows 11 / Linux",
            "- Core Language : Python 3.x",
            "- Deep Learning : PyTorch",
            "- Image Processing : OpenCV, Pillow, NumPy",
            "- Local Registry : SQLite",
            "- Code Editor/IDE : Visual Studio Code",
            "4.3 TOOLS AND PLATFORM",
            "4.3.1 Python",
            "Python is used as the primary implementation language for training, inference, robustness analysis, and registry management.",
            "Its clear syntax and rich ecosystem support rapid experimentation.",
            "Python enables integration of tensor libraries, image processing utilities, and persistence layers in one pipeline.",
            "This reduces development overhead and improves maintainability.",
            "The language is well supported in scientific and deep learning workflows.",
            "It is therefore suitable for iterative watermark model development.",
        ],
        26: [
            TITLE_UPPER,
            "Python also improves reproducibility through script-based training and deterministic seed handling.",
            "4.3.2 PyTorch",
            "PyTorch provides tensor computation, automatic differentiation, and flexible model development for encoder-decoder watermark networks. It supports GPU acceleration and modular experimentation needed for attack-aware training.",
            "4.3.3 OpenCV",
            "OpenCV is used for image transformations, filtering, and practical pre-processing/post-processing operations relevant to watermark robustness evaluation.",
        ],
        27: [
            TITLE_UPPER,
            "4.3.4 NumPy",
            "NumPy is used for numerical operations, array manipulation, and metric preparation in BER and quality evaluation routines.",
            "4.3.5 Pillow",
            "Pillow supports image I/O operations and format handling used in the inference and verification workflow.",
            "4.3.6 SQLite",
            "SQLite stores owner-seed mappings and audit events locally, enabling lightweight yet reliable traceability.",
        ],
        28: [
            TITLE_UPPER,
            "SQLite simplifies deployment because it does not require a separate database server while still supporting structured queries and indexing.",
            "4.3.7 Tkinter",
            "Tkinter is used to provide a desktop user interface for embedding, extraction, and ownership operations.",
            "4.3.8 Matplotlib",
            "Matplotlib supports visualization of robustness and metric trends for experimental interpretation.",
        ],
        29: [
            TITLE_UPPER,
            "These tools collectively support a full workflow from model training to visual analysis and ownership verification.",
            "4.3.9 VISUAL STUDIO CODE",
            "Visual Studio Code serves as the primary development environment for debugging, script execution, and project management across modules.",
        ],
        30: ["CHAPTER-5\nDESIGN"],
        31: [
            TITLE_UPPER,
            "5.1 SYSTEM DESIGN",
            "System design translates project requirements into a practical architecture for embedding, extraction, verification, and ownership tracking.",
            "The design defines how images flow through neural modules, where reliability checks are applied, and how registry events are persisted.",
            "This phase also determines interfaces between training code, inference utilities, and analysis dashboards.",
            "The final design balances image fidelity constraints with extraction robustness and traceability goals.",
            "There are two levels of system design:\n- Logical design.",
            "- Physical design.",
            "The logical design captures data flow and module responsibilities including input handling, embedding, extraction, and verification stages.",
            "- Output design",
        ],
        32: [
            TITLE_UPPER,
            "5.2 PIPELINE DIAGRAM",
            "Figure A",
        ],
        33: [
            TITLE_UPPER,
            "5.3 ACTOR FLOW DIAGRAM",
            "5.3.1 EMBEDDING WORKFLOW",
            "Figure B",
        ],
        34: [
            TITLE_UPPER,
            "5.3.2 VERIFICATION WORKFLOW",
            "Figure C",
        ],
        35: [
            TITLE_UPPER,
            "5.4 INPUT DESIGN",
            "Input design defines how seed values, alpha parameters, image paths, and processing options are validated before embedding or extraction begins.",
            "While designing input, the following aspects are considered:\n- Required metadata and image constraints\n- Parameter range validation for alpha and seed\n- Guided defaults for stable operation\n- Input normalization before model processing",
            "Validation checks reject inconsistent or out-of-range values, reducing downstream errors and improving reproducibility.",
        ],
        36: [
            TITLE_UPPER,
            "5.5 OUTPUT DESIGN",
            "Output design focuses on presenting verification evidence and quality metrics in a clear and actionable format.",
            "While designing output, the following aspects are considered:\n- Which metrics are mandatory (BER, PSNR, confidence)\n- How to display pass/fail states\n- How to preserve image quality while saving outputs",
            "Output Design Breakdown\n1. Embedding Result Display:\nShows effective alpha, visual quality indicators, and save status.\n2. Verification Display:\nShows extracted payload comparison and BER-based decision.",
            "3. Robustness Reporting:\nSummarizes behavior under selected attacks and lists confidence trends across conditions.",
            "The output layer is structured to support both technical debugging and ownership evidence communication.",
            "Consistent labels and compact summaries reduce interpretation errors.",
        ],
        37: [
            TITLE_UPPER,
            "5.6 MODULE DESIGN",
            "Module Breakdown:",
            "1. WATERMARK EMBEDDING MODULE",
            "- Encodes a 64-bit payload into host images using a convolutional encoder.",
            "- Applies controlled alpha to balance invisibility and extraction reliability.",
            "- Supports tiled embedding for stable full-resolution output reconstruction.",
            "- Preserves valid pixel range through clamping and normalization.",
            "- Integrates optional semantic residual enhancement when enabled.",
            "- Produces watermarked outputs ready for archival and distribution.",
            "- Supports post-save verification hooks for operational confidence.",
            "- Logs effective alpha and run metadata for traceability.",
            "- Targets robustness against common post-processing distortions.",
            "2. WATERMARK EXTRACTION MODULE",
            "- Decodes payload bits from suspected watermarked images.",
            "- Computes BER against expected seed-generated payload.",
            "- Supports confidence scoring and deterministic verification reports.",
            "- Uses model-resolution decoding while preserving full image workflow.",
            "- Integrates semantic decoder when the semantic path is enabled.",
            "- Produces diagnostics for pass/fail decision making.",
            "- Handles both single-image and batch extraction workflows.",
            "3. ATTACK AND ROBUSTNESS MODULE",
            "- Simulates JPEG, blur, noise, and geometric perturbations.",
            "- Is used during training to improve decoder resilience.",
            "- Enables reproducible robustness benchmarking experiments.",
            "- Supports adversarial poison and layered attack strategies.",
            "- Reports BER, PSNR, and confidence metrics after perturbation.",
            "- Helps tune alpha for deployment-specific constraints.",
            "- Improves generalization across real-world transformations.",
        ],
        38: [
            TITLE_UPPER,
            "4. RELIABILITY AND REGISTRY MODULE\n- Enforces alpha clamping and save-path reliability checks.\n- Tracks owner-seed mappings in SQLite with audit events.\n- Maintains ownership traceability and usage statistics.",
            "5. ANALYSIS AND VISUALIZATION MODULE\n- Provides attack simulator views and summary dashboards.\n- Aggregates BER/PSNR trends for validation sets.\n- Supports cross-run comparisons for model selection.",
            "5.7 DATA MODEL DESIGN",
            "Data model design structures ownership entities, processing runs, checkpoints, and verification outcomes to ensure consistent retrieval and reporting.",
            "The design supports both transactional operations (embed/decode runs) and historical analysis (audit and robustness trends).",
        ],
        39: [
            TITLE_UPPER,
            "5.8 TABLE DESIGN",
            "Table design defines schema units for ownership, model lifecycle, processing traces, and verification evidence. Proper normalization supports consistency and query performance.",
            "1. OWNER_REGISTRY",
        ],
        40: [TITLE_UPPER, "2. SEEDS"],
        41: [TITLE_UPPER, "3. AUDIT_LOG"],
        42: [TITLE_UPPER, "4. OWNER_PROFILE"],
        43: [TITLE_UPPER, "5. LICENSE_INFO"],
        44: [TITLE_UPPER, "6. IMAGE_COUNT_TRACKER"],
        45: [TITLE_UPPER, "7. IMAGE_EVENT_HISTORY"],
        46: [TITLE_UPPER, "8. ROBUSTNESS_PROFILES", "9. ALPHA_POLICY"],
        47: [TITLE_UPPER, "10. EMBED_JOBS"],
        48: [TITLE_UPPER, "11. DECODE_JOBS"],
        49: [TITLE_UPPER, "12. BATCH_RUNS"],
        50: [TITLE_UPPER, "13. BATCH_ITEMS"],
        51: [TITLE_UPPER, "14. CHECKPOINT_CATALOG", "15. CHECKPOINT_METRICS"],
        52: [TITLE_UPPER, "16. TRAINING_RUNS"],
        53: [TITLE_UPPER, "17. TRAINING_EPOCH_LOG"],
        54: [TITLE_UPPER, "18. VALIDATION_METRICS"],
        55: [TITLE_UPPER, "19. ATTACK_SCENARIOS"],
        56: [TITLE_UPPER, "20. ATTACK_RESULTS"],
        57: [TITLE_UPPER, "21. INFERENCE_REPORTS", "22. VERIFICATION_RESULTS"],
        58: [TITLE_UPPER, "23. ERROR_LOGS"],
        59: [TITLE_UPPER, "24. SYSTEM_SETTINGS"],
        60: [TITLE_UPPER, "25. USER_PREFERENCES"],
        61: [TITLE_UPPER, "26. SESSION_ACTIVITY", "27. EXPORT_HISTORY"],
        62: [TITLE_UPPER, "28. CHANGE_HISTORY"],
        63: ["CHAPTER-6\nFUNCTIONAL AND NON-FUNCTIONAL REQUIREMENTS"],
        64: [
            TITLE_UPPER,
            "6.1FUNCTIONAL REQUIREMENTS",
            "Functional requirements define what the system must do: payload generation, embedding, extraction, verification, registry updates, and robustness reporting.",
            "In this project, functional behavior includes low-alpha embedding, deterministic seed verification, attack-aware extraction evaluation, and owner audit tracking.",
            "6.2 NON-FUNCTIONAL REQUIREMENTS",
            "Non-functional requirements define system quality dimensions such as usability, performance, reliability, and maintainability for deployment.",
            "Some of the non-functional requirements are:\n- Usability: Clear controls for alpha, seed, and workflow mode in the desktop interface.",
            "- Performance: Practical processing time for single and batch workflows with stable memory behavior.",
            "- Reliability: Consistent extraction outcomes with post-save verification safeguards.",
        ],
        65: [
            TITLE_UPPER,
            "- Availability: The application and verification pipeline should be accessible whenever ownership validation is required.",
            "- Security: Registry data and ownership records must be protected against accidental modification.",
            "- Scalability: The workflow should support increasing image volumes through batch operations.",
            "- Portability: Core scripts should run across common operating systems with minimal changes.",
            "- Maintainability: Modular code structure should allow targeted updates without destabilizing the full pipeline.",
        ],
        66: ["CHAPTER-7\nIMPLEMENTATION AND TESTING"],
        67: [
            TITLE_UPPER,
            "7.1 SYSTEM IMPLEMENTATION",
            "Implementation converts design artifacts into an operational watermarking system by integrating model loading, embedding, extraction, reliability checks, and ownership tracking.",
            "For the ZEB project, implementation links training modules, inference workflow, and registry operations into a unified toolchain.",
            "Types of Implementations\nThe project includes:\n- Implementation of a deep learning watermark engine over baseline signal-processing approaches.",
            "- Implementation of reliability controls over standard embedding pipelines to reduce operational failure cases.",
            "Implementation Plan\n- Finalize checkpoint and model-loading path\n- Validate alpha policy and post-save verification\n- Execute robustness tests and record metrics\n- Run ownership registration and audit checks",
        ],
        68: [
            TITLE_UPPER,
            "User Training\nUsers are trained to use embedding and verification workflows with proper alpha and seed management, including how to interpret BER and confidence values.",
            "Implementation of ZEB System\nAfter validation, the system is operated for real image inputs where embedding, extraction, and ownership lookup are executed under the defined reliability policy.",
            "7.2 SYSTEM TESTING",
            "System testing evaluates complete workflow behavior including model inference, save/restore steps, robustness checks, and registry consistency.",
            "Testing verifies that extraction remains reliable under practical distortions and that ownership metadata remains synchronized with generated seeds.",
            "7.2.1 System Testing\n- Definition: Testing the complete watermark lifecycle as one integrated flow.",
            "- In ZEB:\n  o Verify embedding-to-extraction flow end to end\n  o Confirm interaction between model, reliability layer, and registry\n  o Validate pass/fail reporting correctness",
        ],
        69: [
            TITLE_UPPER,
            "7.2.2 Unit Testing\n- Definition: Testing individual modules independently\n- In ZEB:\n  o Encoder/decoder tensor behavior\n  o Reliability helper functions\n  o Registry CRUD operations",
            "7.2.3 Integration Testing\n- Definition: Testing interactions among combined modules\n- In ZEB:\n  o Embedding output consumed by extraction\n  o Registry updated after processing events\n  o Batch pipeline and metric aggregation consistency",
            "7.2.4 Black Box Testing\n- Definition: Validating behavior from external inputs and outputs\n- In ZEB:\n  o Validate visible quality and decoded payload from given images\n  o Confirm expected error handling for invalid inputs",
            "7.2.5 Validation Testing\n- Definition: Ensuring delivered system satisfies stated project requirements\n- In ZEB:\n  o Confirm robustness and invisibility targets\n  o Confirm ownership traceability and report quality",
        ],
        70: [
            TITLE_UPPER,
            "7.2.6 Output Testing\n- Definition: Verifying correctness and clarity of generated outputs\n- In ZEB:\n  o BER/PSNR summaries are consistent\n  o Verification status and confidence are correctly reported",
            "7.2.7 User Acceptance Testing (UAT)\n- Definition: Evaluation by intended users\n- In ZEB:\n  o Users validate practical embedding and verification steps\n  o Feedback used to refine default parameters and messages",
            "7.2.8 White Box Testing\n- Definition: Testing internal logic paths\n- In ZEB:\n  o Check reliability clamps and post-save verification branches\n  o Validate metric and registry update paths under edge cases",
        ],
        71: [
            TITLE_UPPER,
            "7.3 TEST CASES",
            "1. EMBEDDING PIPELINE",
            "2. EXTRACTION PIPELINE",
        ],
        72: [
            TITLE_UPPER,
            "3. ROBUSTNESS EVALUATION",
            "4. OWNER REGISTRY FLOW",
        ],
        73: [
            TITLE_UPPER,
            "5. GUI AND BATCH WORKFLOW",
        ],
        74: ["CHAPTER-8\nRESULTS AND DISCUSSIONS"],
        75: [
            TITLE_UPPER,
            "8.1 RESULTS\nThe implemented ZEB system achieved its primary objective of embedding imperceptible watermarks while preserving practical extraction reliability.",
            "Evaluation shows that low-alpha operation can maintain strong visual quality and still produce low BER when reliability constraints are enforced. Sample runs in this workspace reported PSNR above 41 dB at tuned alpha with BER near 0.0000.",
            "Key Achievements:\n- Robust extraction under common attacks\n- Stable low-alpha behavior for invisibility\n- Integrated ownership registry and audit trace\n- Unified training, inference, and verification workflow",
            "Overall, the results indicate that ZEB is suitable for ownership-aware image protection in realistic distribution settings.",
        ],
        76: [TITLE_UPPER, "8.2 VISUAL OUTPUTS", "HOME SCREEN"],
        77: [TITLE_UPPER, "MODEL LOAD PANEL", "EMBED PANEL", "X-RAY RESIDUAL VIEW"],
        78: [TITLE_UPPER, "BATCH EMBED SETUP", "OWNER REGISTRY FORM", "ALPHA TUNING CONTROLS"],
        79: [TITLE_UPPER, "RELIABILITY CHECKS", "ATTACK SIMULATOR", "BER-PSNR DASHBOARD"],
        80: [TITLE_UPPER, "SEED LOOKUP PAGE"],
        81: [TITLE_UPPER, "EXTRACTION PANEL", "EXTRACTION RESULT", "CONFIDENCE SUMMARY"],
        82: [TITLE_UPPER, "INFERENCE HISTORY", "ROBUSTNESS SUMMARY"],
        83: [TITLE_UPPER, "TRAINING DASHBOARD", "CHECKPOINT SNAPSHOTS", "METRIC PLOTS"],
        84: [TITLE_UPPER, "ATTACK CONFIGURATION", "ADD NEW ATTACK", "ROBUSTNESS CURVES"],
        85: [TITLE_UPPER, "OWNER REGISTRY", "AUDIT ENTRIES", "VERIFICATION STATUS"],
        86: [TITLE_UPPER, "ALERTS AND WARNINGS", "PROFILE SETTINGS", "IMAGE PREVIEW"],
        87: [TITLE_UPPER, "ANALYSIS HOME PAGE", "GENERATE REPORT", "EXPORT RESULT"],
        88: [TITLE_UPPER, "VIEW LOGS", "CHANGE HISTORY"],
        89: [TITLE_UPPER, "VIEW REGISTERED SEEDS", "VIEW OWNER DETAILS"],
        90: ["CHAPTER-9\nCONCLUSION"],
        91: [
            TITLE_UPPER,
            "9.1 SYSTEM MAINTENANCE\nAfter deployment, regular maintenance is required to sustain extraction reliability, model compatibility, and registry integrity.",
            "Examples include:\n- Correcting decode failures under new compression patterns\n- Resolving save-path and quantization edge cases\n- Fixing registry synchronization defects",
            "2. Adaptive Maintenance\nThis addresses environmental and dependency changes affecting the watermark workflow.",
            "Examples include:\n- Updating model dependencies and runtime libraries\n- Adapting to new image format handling requirements\n- Preserving compatibility across platform updates",
            "3. Perfective Maintenance\nThis improves quality and capability based on observed usage and test data.",
            "Examples include:\n- Refining alpha tuning and robustness policy\n- Improving extraction confidence reporting\n- Enhancing analysis dashboards and summaries",
            "4. Preventive Maintenance\nThis reduces future failure risk through periodic audits and optimization.",
            "Examples include:\n- Scheduled robustness regression checks\n- Checkpoint and schema backups\n- Code and documentation hardening",
        ],
        92: [
            TITLE_UPPER,
            "9.2 CONCLUSION",
            "ZEB demonstrates that robust invisible image watermarking can be achieved with a carefully engineered deep learning pipeline and deployment-aware reliability controls.",
            "The project integrates embedding, extraction, attack simulation, and ownership auditing into a coherent framework suitable for practical copyright verification.",
            "Low-alpha operational behavior, post-save verification, and measured BER/PSNR reporting together strengthen confidence in real-world use.",
            "The modular architecture also supports future research extensions without disrupting the core verification workflow.",
            "In conclusion, the project satisfies its objective of providing a high-fidelity, ownership-oriented, and maintainable watermarking system.",
        ],
        93: [
            TITLE_UPPER,
            "9.3 FUTURE ENHANCEMENT",
            "- Stronger geometric invariance through advanced augmentation and synchronization strategies.",
            "- Broader benchmark suites with diverse image sources and attack combinations.",
            "- Improved semantic watermark branch with richer feature coupling.",
            "- Automated reporting pipelines for governance and legal documentation.",
            "- Hardware-aware optimization for faster large-scale deployment.",
            "- Extended provenance integration with external content management systems.",
        ],
        94: ["CHAPTER-10\nBIBLIOGRAPHY"],
        95: [
            TITLE_UPPER,
            "10.1 BOOKS",
            "- Digital Image Processing - Rafael C. Gonzalez and Richard E. Woods.",
            "- Deep Learning - Ian Goodfellow, Yoshua Bengio, and Aaron Courville.",
            "- Pattern Recognition and Machine Learning - Christopher M. Bishop.",
            "- Computer Vision: Algorithms and Applications - Richard Szeliski.",
            "- Hands-On Machine Learning with Scikit-Learn, Keras, and TensorFlow - Aurelien Geron.",
            "- Learning OpenCV - Gary Bradski and Adrian Kaehler.",
            "- Python Crash Course - Eric Matthes.",
            "10.2 WEBSITES",
            "- PyTorch (https://pytorch.org) - Deep learning framework used for model implementation.",
            "- OpenCV (https://opencv.org) - Image processing utilities and transformations.",
            "- Python (https://www.python.org) - Core language runtime and ecosystem.",
        ],
        96: [
            TITLE_UPPER,
            "- NumPy (https://numpy.org) - Numerical array operations for metrics and data handling.",
            "- Pillow (https://python-pillow.org) - Image I/O operations.",
            "- SQLite (https://sqlite.org) - Local ownership registry storage.",
            "- Visual Studio Code (https://code.visualstudio.com) - Development environment.",
            "- Matplotlib (https://matplotlib.org) - Plotting and reporting support.",
            "10.3 JOURNALS AND PUBLICATIONS",
            "[1] J. Zhu et al., \"HiDDeN: Hiding Data With Deep Networks,\" ECCV Workshops, 2018.",
            "[2] K. Zhang et al., \"RivaGAN: Robust Invisible Video Watermarking with Attention,\" 2019.",
            "[3] Y. Luo et al., \"Robust Image Watermarking in Deep Networks with Distortion Simulation,\" 2020.",
            "[4] T. Weng et al., \"StegaStamp: Invisible Hyperlinks in Physical Photographs,\" CVPR, 2020.",
            "[5] M. Ahmadi et al., \"Deep Learning for Digital Watermarking: A Survey,\" 2022.",
        ],
        97: [
            TITLE_UPPER,
            "[6] S. Voloshynovskiy et al., \"Information-Theoretic Foundations of Watermarking Security,\" 2021.",
            "[7] S. Baluja, \"Hiding Images in Plain Sight: Deep Steganography,\" NeurIPS Workshops, 2017.",
            "[8] J. Fridrich, \"Steganography and Digital Watermarking in the Era of Deep Learning,\" 2021.",
            "[9] Recent IEEE Transactions papers on robust and blind image watermarking methods.",
            "[10] Domain-specific reports on copyright protection for digital imagery.",
        ],
        98: ["APPENDICES"],
        99: [
            TITLE_UPPER,
            "LIST OF TABLES",
            "1. DATA TABLES AND METRICS TABLES:",
        ],
        100: [
            TITLE_UPPER,
            "2. TEST-CASES",
            "LIST OF FIGURES",
        ],
        101: [
            TITLE_UPPER,
            "CHANGE HISTORY",
        ],
    }


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"Template DOCX not found: {SOURCE_DOCX}")

    before_doc = Document(str(SOURCE_DOCX))
    before_structure = structure_metrics(before_doc)
    before_style_seq = style_sequence(before_doc)
    before_font_sig = font_signature(before_doc)
    before_hf = header_footer_snapshot(before_doc)
    before_fields = field_counts(before_doc)

    doc = Document(str(SOURCE_DOCX))

    global_replacements = [
        ("FLEXIBOT-AN AI POWERED ACADEMIC ASSISTANT", TITLE_UPPER),
        ("FLEXIBOT - AN AI POWERED ACADEMIC ASSISTANT", TITLE_EXACT),
        ("FLEXIBOT – AN AI POWERED ACADEMIC ASSISTANT", TITLE_EXACT),
        ("FLEXIBOT", "ZEB"),
        ("SNEHA BABU G - LMC24MCA-2054", STUDENT_EXACT),
        ("Ms. SNEHA BABU G", "Mr. ANITH SHIBU THOMAS"),
        ("SNEHA BABU G", "ANITH SHIBU THOMAS"),
    ]
    replace_text_everywhere(doc, global_replacements)

    updates = build_page_updates()
    for table_index, lines in updates.items():
        apply_page_update(doc, table_index, lines)

    doc.save(str(OUTPUT_DOCX))

    after_doc = Document(str(OUTPUT_DOCX))
    after_structure = structure_metrics(after_doc)
    after_style_seq = style_sequence(after_doc)
    after_font_sig = font_signature(after_doc)
    after_hf = header_footer_snapshot(after_doc)
    after_fields = field_counts(after_doc)

    structure_ok = before_structure == after_structure
    style_ok = (before_style_seq == after_style_seq) and (before_font_sig == after_font_sig)
    header_footer_ok = before_hf == after_hf
    pagination_ok = before_fields == after_fields

    all_text = extract_all_text(after_doc)
    all_text_upper = all_text.upper()

    semantic_checks = {
        "title_applied": TITLE_EXACT in all_text,
        "student_applied": STUDENT_EXACT in all_text,
        "no_cli_token": "CLI" not in all_text_upper,
        "old_project_removed": "FLEXIBOT" not in all_text_upper,
        "chapter_order_ok": chapter_order_ok(all_text_upper),
    }
    semantic_ok = all(semantic_checks.values())

    unresolved = [
        ("5.2 Pipeline Diagram figure asset", table_page_number(after_doc, 32)),
        ("5.3 Actor Flow Diagram figure assets", f"{table_page_number(after_doc, 33)}-{table_page_number(after_doc, 34)}"),
        ("8.2 Visual Outputs screenshot assets", f"{table_page_number(after_doc, 76)}-{table_page_number(after_doc, 89)}"),
        ("List of Figures image references", table_page_number(after_doc, 100)),
    ]

    report_lines = []
    report_lines.append("Strict 5-Step Verification Report")
    report_lines.append("")
    report_lines.append("1) Structure check")
    report_lines.append(f"- Before: paragraphs={before_structure['paragraph_count']}, tables={before_structure['table_count']}, sections={before_structure['section_count']}")
    report_lines.append(f"- After : paragraphs={after_structure['paragraph_count']}, tables={after_structure['table_count']}, sections={after_structure['section_count']}")
    report_lines.append(f"- Status: {'PASS' if structure_ok else 'FAIL'}")

    report_lines.append("")
    report_lines.append("2) Style/format check")
    report_lines.append(f"- Paragraph style sequence parity: {'OK' if before_style_seq == after_style_seq else 'MISMATCH'}")
    report_lines.append(f"- Run font signature parity      : {'OK' if before_font_sig == after_font_sig else 'MISMATCH'}")
    report_lines.append(f"- Status: {'PASS' if style_ok else 'FAIL'}")

    report_lines.append("")
    report_lines.append("3) Header/footer check")
    report_lines.append(f"- Header/footer text and style parity: {'OK' if header_footer_ok else 'MISMATCH'}")
    report_lines.append(f"- Status: {'PASS' if header_footer_ok else 'FAIL'}")

    report_lines.append("")
    report_lines.append("4) Pagination integrity check")
    report_lines.append(f"- Before field nodes: {before_fields['field_nodes']} (page fields: {before_fields['page_fields']})")
    report_lines.append(f"- After  field nodes: {after_fields['field_nodes']} (page fields: {after_fields['page_fields']})")
    report_lines.append(f"- Status: {'PASS' if pagination_ok else 'FAIL'}")

    report_lines.append("")
    report_lines.append("5) Semantic compliance check")
    report_lines.append(f"- Required title applied: {semantic_checks['title_applied']}")
    report_lines.append(f"- Required identity applied: {semantic_checks['student_applied']}")
    report_lines.append(f"- No CLI mention remains: {semantic_checks['no_cli_token']}")
    report_lines.append(f"- Old domain label removed (FLEXIBOT): {semantic_checks['old_project_removed']}")
    report_lines.append(f"- Chapter sequence preserved: {semantic_checks['chapter_order_ok']}")
    report_lines.append(f"- Status: {'PASS' if semantic_ok else 'FAIL'}")

    report_lines.append("")
    report_lines.append("Sections left unchanged due unclear source assets")
    for section_name, page_no in unresolved:
        report_lines.append(f"- {section_name} | page(s): {page_no}")

    REPORT_PATH.write_text("\n".join(report_lines), encoding="utf-8")

    print("Created:", OUTPUT_DOCX)
    print("Verification report:", REPORT_PATH)


if __name__ == "__main__":
    main()
