from __future__ import annotations

from collections import Counter
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, List, Tuple

from docx import Document

BASE = Path("reference_template.docx")
OUT = Path("anith.docx")
REPORT = Path("analysis/recheck_report.txt")
JSON_OUT = Path("analysis/recheck_parity.json")


@dataclass
class SectionProps:
    start_type: str
    orientation: str
    page_width: int
    page_height: int
    left_margin: int
    right_margin: int
    top_margin: int
    bottom_margin: int
    header_distance: int
    footer_distance: int
    different_first_page_header_footer: bool


def sec_props(section) -> SectionProps:
    return SectionProps(
        start_type=str(section.start_type),
        orientation=str(section.orientation),
        page_width=int(section.page_width),
        page_height=int(section.page_height),
        left_margin=int(section.left_margin),
        right_margin=int(section.right_margin),
        top_margin=int(section.top_margin),
        bottom_margin=int(section.bottom_margin),
        header_distance=int(section.header_distance),
        footer_distance=int(section.footer_distance),
        different_first_page_header_footer=bool(section.different_first_page_header_footer),
    )


def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def style_sequence(doc: Document) -> List[str]:
    seq: List[str] = []
    for p in iter_all_paragraphs(doc):
        seq.append(p.style.name if p.style is not None else "")
    return seq


def para_format_signature(doc: Document) -> Counter:
    c = Counter()
    for p in iter_all_paragraphs(doc):
        fmt = p.paragraph_format
        key = (
            p.style.name if p.style is not None else "",
            int(fmt.left_indent) if fmt.left_indent is not None else None,
            int(fmt.right_indent) if fmt.right_indent is not None else None,
            int(fmt.first_line_indent) if fmt.first_line_indent is not None else None,
            float(fmt.line_spacing) if isinstance(fmt.line_spacing, (int, float)) else None,
            str(fmt.line_spacing_rule) if fmt.line_spacing_rule is not None else None,
            int(fmt.space_before) if fmt.space_before is not None else None,
            int(fmt.space_after) if fmt.space_after is not None else None,
            str(fmt.alignment) if fmt.alignment is not None else None,
        )
        c[key] += 1
    return c


def run_font_signature(doc: Document) -> Counter:
    c = Counter()
    for p in iter_all_paragraphs(doc):
        for r in p.runs:
            key = (
                r.font.name,
                float(r.font.size.pt) if r.font.size is not None else None,
                bool(r.bold),
                bool(r.italic),
                bool(r.underline),
            )
            c[key] += 1
    return c


def table_shape_signature(doc: Document) -> List[Tuple[int, int]]:
    sig: List[Tuple[int, int]] = []
    for t in doc.tables:
        rows = len(t.rows)
        cols = len(t.columns)
        sig.append((rows, cols))
    return sig


def header_footer_setup_signature(doc: Document) -> List[Dict[str, object]]:
    out: List[Dict[str, object]] = []
    for i, s in enumerate(doc.sections):
        out.append(
            {
                "section": i,
                "is_linked_header": bool(s.header.is_linked_to_previous),
                "is_linked_footer": bool(s.footer.is_linked_to_previous),
                "header_para_count": len(s.header.paragraphs),
                "footer_para_count": len(s.footer.paragraphs),
                "different_first": bool(s.different_first_page_header_footer),
                "odd_even_pages": bool(doc.settings.odd_and_even_pages_header_footer),
            }
        )
    return out


def field_counts(doc: Document) -> Dict[str, int]:
    parts = [doc.part]
    seen = {id(doc.part)}
    for s in doc.sections:
        for part in (s.header.part, s.footer.part):
            pid = id(part)
            if pid not in seen:
                seen.add(pid)
                parts.append(part)

    field_nodes = 0
    page_field_nodes = 0
    for part in parts:
        root = part.element
        fld_simple = root.xpath('.//*[local-name()="fldSimple"]')
        instr = root.xpath('.//*[local-name()="instrText"]')
        field_nodes += len(fld_simple) + len(instr)

        for n in fld_simple:
            instr_attr = n.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr") or ""
            if "PAGE" in instr_attr.upper():
                page_field_nodes += 1

        for n in instr:
            txt = n.text or ""
            if "PAGE" in txt.upper():
                page_field_nodes += 1

    return {"field_nodes": field_nodes, "page_field_nodes": page_field_nodes}


def compare_docs(base_path: Path, out_path: Path) -> Dict[str, object]:
    b = Document(str(base_path))
    o = Document(str(out_path))

    base_sections = [sec_props(s) for s in b.sections]
    out_sections = [sec_props(s) for s in o.sections]

    data = {
        "paragraph_count_base": len(b.paragraphs),
        "paragraph_count_out": len(o.paragraphs),
        "table_count_base": len(b.tables),
        "table_count_out": len(o.tables),
        "section_count_base": len(b.sections),
        "section_count_out": len(o.sections),
        "sections_equal": [asdict(x) for x in base_sections] == [asdict(x) for x in out_sections],
        "style_sequence_equal": style_sequence(b) == style_sequence(o),
        "paragraph_format_signature_equal": para_format_signature(b) == para_format_signature(o),
        "run_font_signature_equal": run_font_signature(b) == run_font_signature(o),
        "table_shape_equal": table_shape_signature(b) == table_shape_signature(o),
        "header_footer_setup_equal": header_footer_setup_signature(b) == header_footer_setup_signature(o),
        "field_counts_base": field_counts(b),
        "field_counts_out": field_counts(o),
        "field_counts_equal": field_counts(b) == field_counts(o),
    }

    return data


def main() -> None:
    if not BASE.exists() or not OUT.exists():
        raise FileNotFoundError("Expected reference_template.docx and anith.docx in workspace root")

    result = compare_docs(BASE, OUT)

    import json

    JSON_OUT.write_text(json.dumps(result, indent=2), encoding="utf-8")

    lines = []
    lines.append("DOCX Parity Recheck")
    lines.append("")
    lines.append(f"Paragraphs: {result['paragraph_count_base']} -> {result['paragraph_count_out']}")
    lines.append(f"Tables: {result['table_count_base']} -> {result['table_count_out']}")
    lines.append(f"Sections: {result['section_count_base']} -> {result['section_count_out']}")
    lines.append(f"Section page setup equal: {result['sections_equal']}")
    lines.append(f"Paragraph style sequence equal: {result['style_sequence_equal']}")
    lines.append(f"Paragraph format signature equal: {result['paragraph_format_signature_equal']}")
    lines.append(f"Run font signature equal: {result['run_font_signature_equal']}")
    lines.append(f"Table shape equal: {result['table_shape_equal']}")
    lines.append(f"Header/footer setup equal: {result['header_footer_setup_equal']}")
    lines.append(f"Field counts equal: {result['field_counts_equal']}")
    lines.append(f"Field counts (base): {result['field_counts_base']}")
    lines.append(f"Field counts (out) : {result['field_counts_out']}")

    REPORT.write_text("\n".join(lines), encoding="utf-8")
    print("WROTE", REPORT)
    print("WROTE", JSON_OUT)


if __name__ == "__main__":
    main()
